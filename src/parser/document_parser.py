import pymupdf  # fitz
from docx import Document
from pathlib import Path
from typing import Dict, Any, Tuple
import re

class DocumentParser:
    """Handles text extraction from PDF and DOCX resumes."""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx"}

    @staticmethod
    def validate_file(file_path: str | Path) -> None:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")
        if path.suffix.lower() not in DocumentParser.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {path.suffix}. Only PDF and DOCX allowed.")

    @staticmethod
    def extract_text_from_pdf(file_path: str | Path) -> Tuple[str, Dict[str, Any]]:
        DocumentParser.validate_file(file_path)
        path = Path(file_path)
        
        metadata = {
            "filename": path.name,
            "file_type": "pdf",
            "page_count": 0,
            "extraction_method": "pymupdf_text"
        }
        
        text_blocks = []
        with pymupdf.open(path) as doc:
            metadata["page_count"] = len(doc)
            for page_num, page in enumerate(doc, 1):
                page_text = page.get_text("text")
                page_text = DocumentParser._clean_text(page_text)
                if page_text.strip():
                    text_blocks.append(f"--- Page {page_num} ---\n{page_text}")
        
        full_text = "\n\n".join(text_blocks)
        return full_text.strip(), metadata

    @staticmethod
    def extract_text_from_docx(file_path: str | Path) -> Tuple[str, Dict[str, Any]]:
        DocumentParser.validate_file(file_path)
        path = Path(file_path)
        
        doc = Document(path)
        metadata = {
            "filename": path.name,
            "file_type": "docx",
            "paragraph_count": len(doc.paragraphs),
            "extraction_method": "python-docx"
        }
        
        paragraphs = [DocumentParser._clean_text(para.text) for para in doc.paragraphs if para.text.strip()]
        full_text = "\n\n".join(paragraphs)
        return full_text.strip(), metadata

    @staticmethod
    def _clean_text(text: str) -> str:
        if not text:
            return ""
        text = re.sub(r'\s+', ' ', text).strip()
        text = re.sub(r'[\u200b\u200c\u200d]', '', text)
        return text

    @staticmethod
    def parse_resume(file_path: str | Path) -> Tuple[str, Dict[str, Any]]:
        path = Path(file_path)
        suffix = path.suffix.lower()
        
        if suffix == ".pdf":
            return DocumentParser.extract_text_from_pdf(file_path)
        elif suffix == ".docx":
            return DocumentParser.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file: {suffix}")

if __name__ == "__main__":
    # Test with your sample resume (place file in data/ folder)
    test_file = "data/sample_resume.pdf"   # Change filename as needed
    try:
        text, meta = DocumentParser.parse_resume(test_file)
        print("Metadata:", meta)
        print("\nExtracted Text (first 500 chars):")
        print(text[:500] + "..." if len(text) > 500 else text)
    except Exception as e:
        print("Error:", e)